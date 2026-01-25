import os
import re
from docx import Document
from docx.shared import Pt
import sys

def parse_markdown_to_docx(md_content, output_path):
    doc = Document()
    
    # Simple style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            # Handle bold in list items
            apply_formatting(p)
        elif line.startswith('1. '):
            p = doc.add_paragraph(line[3:], style='List Number')
             # Handle bold in list items
            apply_formatting(p)
        elif line.startswith('---'):
            doc.add_page_break()
        else:
            p = doc.add_paragraph(line)
            apply_formatting(p)

    doc.save(output_path)
    print(f"File saved to: {output_path}")

def apply_formatting(paragraph):
    # Rudimentary bold handling: **text**
    # This is tricky in python-docx after creating paragraph, 
    # but let's try a simple replace approach if needed or just leave as is for now 
    # as robust MD parsing is complex.
    # However, for a presentation file, headings and lists are most important.
    pass

if __name__ == "__main__":
    input_file = r"c:\Users\Tommaso\Documents\HOWEST\TeamProject\MTS3-MCTE-Team-Project-Energy-G1\docs\deliverables\06_PROJECT_MANAGEMENT.md"
    output_file = r"c:\Users\Tommaso\Documents\HOWEST\TeamProject\MTS3-MCTE-Team-Project-Energy-G1\docs\deliverables\06_PROJECT_MANAGEMENT_FINAL.docx"
    
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
        
    parse_markdown_to_docx(content, output_file)
